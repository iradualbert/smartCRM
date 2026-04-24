from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DIR = ROOT / "corebackend" / "sales" / "default"
FILES_DIR = ROOT / "corebackend" / "files" / "templates"

NAVY = RGBColor(15, 23, 42)
SLATE = RGBColor(71, 85, 105)
MUTED = RGBColor(100, 116, 139)
BORDER = "D6DFEA"
HEADER_FILL = "EAF1F8"
PANEL_FILL = "F8FBFE"
TOTAL_FILL = "EDF8F1"


def set_cell_background(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size=8):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_table_borders(table, color=BORDER, size=8):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in("w:tblBorders")
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tbl_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tbl_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), color)


def set_cell_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(inches * 1440)))


def set_table_layout(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_run(paragraph, text, *, size=10.5, bold=False, color=SLATE):
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return run


def add_text_line(cell, text, *, size=10.5, bold=False, color=SLATE, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.add_paragraph() if cell.text or len(cell.paragraphs) > 1 or cell.paragraphs[0].text else cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(0)
    add_run(paragraph, text, size=size, bold=bold, color=color)
    return paragraph


def add_label_and_value(cell, label, value):
    clear_cell(cell)
    add_text_line(cell, label.upper(), size=8, bold=True, color=MUTED)
    add_text_line(cell, value, size=12, bold=True, color=NAVY)


def finalize(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style.font.size = Pt(10.5)
    style.font.color.rgb = SLATE


def spacer(doc, points=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def add_header(doc, title):
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo.paragraph_format.space_after = Pt(4)
    add_run(logo, "[[CompanyLogo]]", size=10, color=MUTED)

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_paragraph.paragraph_format.space_after = Pt(0)
    add_run(title_paragraph, title.upper(), size=22, bold=True, color=NAVY)

    number = doc.add_paragraph()
    number.alignment = WD_ALIGN_PARAGRAPH.LEFT
    number.paragraph_format.space_after = Pt(10)
    add_run(number, "[[DocumentNumber]]", size=11.5, bold=True, color=SLATE)


def add_party_block(doc, customer_heading):
    table = doc.add_table(rows=1, cols=2)
    set_table_layout(table, [3.1, 3.1])
    set_table_borders(table)
    table.cell(0, 0).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_background(table.cell(0, 0), PANEL_FILL)
    set_cell_background(table.cell(0, 1), PANEL_FILL)
    set_cell_borders(table.cell(0, 0))
    set_cell_borders(table.cell(0, 1))

    company = table.cell(0, 0)
    customer = table.cell(0, 1)

    clear_cell(company)
    add_text_line(company, "FROM", size=8, bold=True, color=MUTED)
    add_text_line(company, "[[CompanyLegalName]]", size=12, bold=True, color=NAVY)
    add_text_line(company, "[[CompanyAddress]]")
    add_text_line(company, "[[CompanyEmail]]")
    add_text_line(company, "[[CompanyPhone]]")
    add_text_line(company, "[[CompanyWebsite]]")
    add_text_line(company, "Tax number: [[CompanyTaxNumber]]")

    clear_cell(customer)
    add_text_line(customer, customer_heading.upper(), size=8, bold=True, color=MUTED)
    add_text_line(customer, "[[ClientName]]", size=12, bold=True, color=NAVY)
    add_text_line(customer, "[[ClientDetails]]")


def add_metadata(doc, fields):
    table = doc.add_table(rows=2, cols=len(fields))
    set_table_layout(table, [6.2 / len(fields)] * len(fields))
    set_table_borders(table)
    for idx, (label, value) in enumerate(fields):
        head = table.cell(0, idx)
        body = table.cell(1, idx)
        set_cell_background(head, HEADER_FILL)
        set_cell_borders(head)
        set_cell_borders(body)
        add_label_and_value(head, label, "")
        clear_cell(body)
        add_text_line(body, value, size=10.5, bold=True, color=NAVY)


def add_lines_table(doc, include_prices=True):
    cols = 4 if include_prices else 2
    widths = [2.9, 0.7, 1.2, 1.2] if include_prices else [4.9, 1.3]
    table = doc.add_table(rows=2, cols=cols)
    set_table_layout(table, widths)
    set_table_borders(table)

    headers = ["Description", "Qty"]
    placeholders = ["[[LineDescription]]", "[[LineQty]]"]
    if include_prices:
        headers += ["Unit price", "Line total"]
        placeholders += ["[[LineUnitPrice]]", "[[LineTotal]]"]

    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_background(cell, HEADER_FILL)
        set_cell_borders(cell)
        clear_cell(cell)
        add_text_line(
            cell,
            header,
            size=9,
            bold=True,
            color=NAVY,
            align=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
        )

    for idx, placeholder in enumerate(placeholders):
        cell = table.cell(1, idx)
        set_cell_borders(cell)
        clear_cell(cell)
        add_text_line(
            cell,
            placeholder,
            size=10.5,
            color=SLATE,
            align=WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER,
        )


def add_totals(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_layout(table, [1.8, 1.35])
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_borders(table)

    for idx, (label, value, emphasize) in enumerate(rows):
        left = table.cell(idx, 0)
        right = table.cell(idx, 1)
        if emphasize:
            set_cell_background(left, TOTAL_FILL)
            set_cell_background(right, TOTAL_FILL)
        set_cell_borders(left)
        set_cell_borders(right)
        clear_cell(left)
        clear_cell(right)
        add_text_line(left, label, size=9 if not emphasize else 9.5, bold=True, color=MUTED if not emphasize else NAVY)
        add_text_line(right, value, size=11 if not emphasize else 12.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.RIGHT)


def add_notes(doc, heading):
    table = doc.add_table(rows=2, cols=1)
    set_table_layout(table, [6.2])
    set_table_borders(table)
    set_cell_background(table.cell(0, 0), HEADER_FILL)
    clear_cell(table.cell(0, 0))
    clear_cell(table.cell(1, 0))
    add_text_line(table.cell(0, 0), heading.upper(), size=8.5, bold=True, color=MUTED)
    add_text_line(table.cell(1, 0), "[[Notes]]", size=10.5, color=SLATE)


def build_invoice():
    doc = Document()
    finalize(doc)
    add_header(doc, "Invoice")
    add_party_block(doc, "Bill to")
    spacer(doc)
    add_metadata(doc, [("Issued", "[[DocumentDate]]"), ("Due", "[[DueDate]]"), ("Currency", "[[Currency]]")])
    spacer(doc)
    add_lines_table(doc, include_prices=True)
    spacer(doc)
    add_totals(doc, [("Subtotal", "[[SubTotal]]", False), ("Tax ([[TaxRatePercent]])", "[[Tax]]", False), ("Total", "[[Total]]", True)])
    spacer(doc)
    add_notes(doc, "Notes")
    return doc


def build_quotation():
    doc = Document()
    finalize(doc)
    add_header(doc, "Quotation")
    add_party_block(doc, "Prepared for")
    spacer(doc)
    add_metadata(doc, [("Date", "[[DocumentDate]]"), ("Valid until", "[[ValidUntil]]"), ("Currency", "[[Currency]]")])
    spacer(doc)
    add_lines_table(doc, include_prices=True)
    spacer(doc)
    add_totals(doc, [("Subtotal", "[[SubTotal]]", False), ("Discount", "[[Discount]]", False), ("Tax ([[TaxRatePercent]])", "[[Tax]]", False), ("Total", "[[Total]]", True)])
    spacer(doc)
    add_notes(doc, "Scope and notes")
    return doc


def build_proforma():
    doc = Document()
    finalize(doc)
    add_header(doc, "Proforma Invoice")
    add_party_block(doc, "Prepared for")
    spacer(doc)
    add_metadata(doc, [("Date", "[[DocumentDate]]"), ("Due", "[[DueDate]]"), ("Currency", "[[Currency]]")])
    spacer(doc)
    add_lines_table(doc, include_prices=True)
    spacer(doc)
    add_totals(doc, [("Subtotal", "[[SubTotal]]", False), ("Discount", "[[Discount]]", False), ("Tax ([[TaxRatePercent]])", "[[Tax]]", False), ("Total", "[[Total]]", True)])
    spacer(doc)
    add_notes(doc, "Notes")
    return doc


def build_receipt():
    doc = Document()
    finalize(doc)
    add_header(doc, "Receipt")
    add_party_block(doc, "Received from")
    spacer(doc)
    add_metadata(doc, [("Date", "[[DocumentDate]]"), ("Invoice", "[[InvoiceNumber]]"), ("Currency", "[[Currency]]")])
    spacer(doc)
    add_totals(doc, [("Amount received", "[[AmountPaid]]", True), ("Payment method", "[[PaymentMethod]]", False)])
    spacer(doc)
    add_notes(doc, "Notes")
    return doc


def build_delivery_note():
    doc = Document()
    finalize(doc)
    add_header(doc, "Delivery Note")
    add_party_block(doc, "Deliver to")
    spacer(doc)
    add_metadata(doc, [("Issued", "[[DocumentDate]]"), ("Delivery date", "[[DeliveryDate]]")])
    spacer(doc)
    add_lines_table(doc, include_prices=False)
    spacer(doc)
    add_notes(doc, "Delivery notes")
    return doc


BUILDERS = {
    "quotation_standard_template.docx": build_quotation,
    "proforma_standard_template.docx": build_proforma,
    "invoice_standard_template.docx": build_invoice,
    "receipt_standard_template.docx": build_receipt,
    "delivery_note_standard_template.docx": build_delivery_note,
}

FILES_BUILDERS = {
    "quotation_pro.docx": build_quotation,
    "proforma_pro.docx": build_proforma,
    "invoice_standard_template.docx": build_invoice,
    "receipt_pro.docx": build_receipt,
    "delivery_note_pro.docx": build_delivery_note,
}


def save_all():
    DEFAULT_DIR.mkdir(parents=True, exist_ok=True)
    FILES_DIR.mkdir(parents=True, exist_ok=True)

    for filename, builder in BUILDERS.items():
        builder().save(str(DEFAULT_DIR / filename))

    for filename, builder in FILES_BUILDERS.items():
        builder().save(str(FILES_DIR / filename))


if __name__ == "__main__":
    save_all()
