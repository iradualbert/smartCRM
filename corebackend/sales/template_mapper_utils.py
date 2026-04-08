from __future__ import annotations

import copy
import platform
import re
import shutil
import subprocess
import tempfile
import time
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.text.paragraph import Paragraph


PLACEHOLDER_PATTERN = re.compile(r"\[\[([A-Za-z0-9_]+)\]\]")


DOCUMENT_SCHEMAS: Dict[str, Dict[str, Any]] = {
    "invoice": {
        "label": "Invoice",
        "suggested_mapping": {
            "DocumentNumber": "document.number",
            "DocumentDate": "document.date",
            "DueDate": "document.due_date",
            "CompanyName": "company.name",
            "CompanyAddress": "company.address",
            "CompanyEmail": "company.email",
            "CompanyPhone": "company.phone",
            "ClientName": "client.name",
            "ClientDetails": "client.details",
            "Currency": "document.currency",
            "SubTotal": "totals.subtotal",
            "Discount": "totals.discount",
            "Tax": "totals.tax",
            "Total": "totals.total",
            "TaxRatePercent": "totals.tax_rate_percent",
            "Notes": "notes",
            "LineDescription": "lines.description",
            "LineQty": "lines.qty",
            "LineUnitPrice": "lines.unit_price",
            "LineTotal": "lines.total",
        },
        "aliases": {
            "InvoiceNumber": "DocumentNumber",
            "InvoiceDate": "DocumentDate",
            "ItemDescription": "LineDescription",
            "ItemQty": "LineQty",
            "ItemUnitPrice": "LineUnitPrice",
            "ItemTotal": "LineTotal",
        },
    },
    "quotation": {
        "label": "Quotation",
        "suggested_mapping": {
            "DocumentNumber": "document.number",
            "DocumentDate": "document.date",
            "ValidUntil": "document.valid_until",
            "CompanyName": "company.name",
            "CompanyAddress": "company.address",
            "CompanyEmail": "company.email",
            "CompanyPhone": "company.phone",
            "ClientName": "client.name",
            "ClientDetails": "client.details",
            "Currency": "document.currency",
            "SubTotal": "totals.subtotal",
            "Discount": "totals.discount",
            "Tax": "totals.tax",
            "Total": "totals.total",
            "TaxRatePercent": "totals.tax_rate_percent",
            "Notes": "notes",
            "LineDescription": "lines.description",
            "LineQty": "lines.qty",
            "LineUnitPrice": "lines.unit_price",
            "LineTotal": "lines.total",
        },
        "aliases": {
            "QuotationNumber": "DocumentNumber",
            "QuotationDate": "DocumentDate",
            "ItemDescription": "LineDescription",
            "ItemQty": "LineQty",
            "ItemUnitPrice": "LineUnitPrice",
            "ItemTotal": "LineTotal",
        },
    },
    "proforma": {
        "label": "Proforma Invoice",
        "suggested_mapping": {
            "DocumentNumber": "document.number",
            "DocumentDate": "document.date",
            "DueDate": "document.due_date",
            "CompanyName": "company.name",
            "CompanyAddress": "company.address",
            "CompanyEmail": "company.email",
            "CompanyPhone": "company.phone",
            "ClientName": "client.name",
            "ClientDetails": "client.details",
            "Currency": "document.currency",
            "SubTotal": "totals.subtotal",
            "Discount": "totals.discount",
            "Tax": "totals.tax",
            "Total": "totals.total",
            "TaxRatePercent": "totals.tax_rate_percent",
            "Notes": "notes",
            "LineDescription": "lines.description",
            "LineQty": "lines.qty",
            "LineUnitPrice": "lines.unit_price",
            "LineTotal": "lines.total",
        },
        "aliases": {
            "ProformaNumber": "DocumentNumber",
            "ProformaDate": "DocumentDate",
            "ItemDescription": "LineDescription",
            "ItemQty": "LineQty",
            "ItemUnitPrice": "LineUnitPrice",
            "ItemTotal": "LineTotal",
        },
    },
    "delivery_note": {
        "label": "Delivery Note",
        "suggested_mapping": {
            "DocumentNumber": "document.number",
            "DocumentDate": "document.date",
            "DeliveryDate": "document.delivery_date",
            "CompanyName": "company.name",
            "CompanyAddress": "company.address",
            "CompanyEmail": "company.email",
            "CompanyPhone": "company.phone",
            "ClientName": "client.name",
            "ClientDetails": "client.details",
            "Warehouse": "document.warehouse",
            "DriverName": "document.driver_name",
            "VehicleNumber": "document.vehicle_number",
            "Notes": "notes",
            "LineDescription": "lines.description",
            "LineQty": "lines.qty",
        },
        "aliases": {
            "DeliveryNoteNumber": "DocumentNumber",
            "ItemDescription": "LineDescription",
            "ItemQty": "LineQty",
        },
    },
    
    "receipt": {
        "label": "Receipt",
        "suggested_mapping": {
            "DocumentNumber": "document.number",
            "DocumentDate": "document.date",
            "CompanyName": "company.name",
            "CompanyAddress": "company.address",
            "CompanyEmail": "company.email",
            "CompanyPhone": "company.phone",
            "ClientName": "client.name",
            "ClientDetails": "client.details",
            "Currency": "document.currency",
            "SubTotal": "totals.subtotal",
            "Discount": "totals.discount",
            "Tax": "totals.tax",
            "Total": "totals.total",
            "TaxRatePercent": "totals.tax_rate_percent",
            "Notes": "notes",
            "AmountPaid": "totals.total",
            "PaymentDate": "document.date",
            "PaymentMethod": "custom.payment_method",
            "InvoiceNumber": "custom.invoice_number",
        },
        "aliases": {
            "ReceiptNumber": "DocumentNumber",
            "ReceiptDate": "DocumentDate",
        },
    },
}


def money(value: Any, currency_symbol: str = "$") -> str:
    amount = Decimal(str(value or 0)).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{currency_symbol}{amount:,.2f}"


def get_document_schema(document_type: str) -> Dict[str, Any]:
    if document_type not in DOCUMENT_SCHEMAS:
        supported = ", ".join(sorted(DOCUMENT_SCHEMAS))
        raise ValueError(f"Unsupported document_type '{document_type}'. Supported: {supported}")
    return DOCUMENT_SCHEMAS[document_type]


def normalize_placeholder_name(placeholder: str, document_type: str) -> str:
    schema = get_document_schema(document_type)
    aliases = schema.get("aliases", {})
    return aliases.get(placeholder, placeholder)


def extract_placeholders_from_docx(template_path: str | Path) -> List[str]:
    doc = Document(str(template_path))
    placeholders = set()

    for para in doc.paragraphs:
        placeholders.update(PLACEHOLDER_PATTERN.findall(para.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                placeholders.update(PLACEHOLDER_PATTERN.findall(cell.text))

    return sorted(placeholders)


def inspect_template(template_path: str | Path, document_type: str = "invoice") -> Dict[str, Any]:
    placeholders = extract_placeholders_from_docx(template_path)
    schema = get_document_schema(document_type)
    suggested = copy.deepcopy(schema["suggested_mapping"])

    normalized_pairs = []
    unmapped = []

    for placeholder in placeholders:
        normalized = normalize_placeholder_name(placeholder, document_type)
        normalized_pairs.append({"raw": placeholder, "normalized": normalized})
        if normalized not in suggested:
            unmapped.append(placeholder)

    line_row_placeholders = sorted(
        p for p in placeholders
        if normalize_placeholder_name(p, document_type)
        in {"LineDescription", "LineQty", "LineUnitPrice", "LineTotal"}
    )

    return {
        "document_type": document_type,
        "document_label": schema["label"],
        "all_placeholders": placeholders,
        "normalized_placeholders": normalized_pairs,
        "suggested_mapping": suggested,
        "unmapped_by_default": unmapped,
        "detected_line_placeholders": line_row_placeholders,
    }


def build_standard_sales_document_data(data: Dict[str, Any], document_type: str = "invoice") -> Dict[str, Any]:
    company = data.get("company", {})
    client = data.get("client", {})
    document = data.get("document", {})
    lines = data.get("lines") or data.get("items") or []

    currency_symbol = document.get("currency_symbol", data.get("currency_symbol", "$"))
    currency_code = document.get("currency", data.get("currency", "USD"))

    normalized_lines = []
    subtotal_raw = Decimal("0.00")

    for line in lines:
        qty = Decimal(str(line.get("qty", 0)))
        unit_price = Decimal(str(line.get("unit_price", 0)))
        line_total_raw = Decimal(str(line.get("total", qty * unit_price)))
        subtotal_raw += line_total_raw

        normalized_lines.append(
            {
                "description": line.get("description", ""),
                "qty": str(qty.to_integral()) if qty == qty.to_integral() else str(qty.normalize()),
                "unit_price": money(unit_price, currency_symbol),
                "total": money(line_total_raw, currency_symbol),
            }
        )

    discount_raw = Decimal(str(data.get("discount", document.get("discount", 0)) or 0))
    tax_rate_raw = Decimal(str(data.get("tax_rate", document.get("tax_rate", 0)) or 0))
    tax_raw = Decimal(str(data.get("tax", document.get("tax", 0)) or 0))
    if tax_raw == 0 and tax_rate_raw != 0:
        tax_raw = (subtotal_raw - discount_raw) * tax_rate_raw

    total_raw = subtotal_raw - discount_raw + tax_raw

    return {
        "document": {
            "number": document.get("number", data.get("document_number", data.get("invoice_number", ""))),
            "date": document.get("date", data.get("document_date", data.get("invoice_date", ""))),
            "due_date": document.get("due_date", data.get("due_date", "")),
            "valid_until": document.get("valid_until", data.get("valid_until", "")),
            "delivery_date": document.get("delivery_date", data.get("delivery_date", "")),
            "warehouse": document.get("warehouse", ""),
            "driver_name": document.get("driver_name", ""),
            "vehicle_number": document.get("vehicle_number", ""),
            "currency": currency_code,
            "currency_symbol": currency_symbol,
        },
        "company": {
            "name": company.get("name", ""),
            "address": company.get("address", ""),
            "email": company.get("email", ""),
            "phone": company.get("phone", ""),
        },
        "client": {
            "name": client.get("name", ""),
            "details": client.get("details", client.get("address", "")),
        },
        "totals": {
            "subtotal": money(subtotal_raw, currency_symbol),
            "discount": money(discount_raw, currency_symbol),
            "tax": money(tax_raw, currency_symbol),
            "total": money(total_raw, currency_symbol),
            "tax_rate_percent": f"{(tax_rate_raw * 100).quantize(Decimal('0.01'))}%",
        },
        "lines": normalized_lines,
        "notes": data.get("notes", ""),
        "document_type": document_type,
    }


def get_value_from_path(data: Dict[str, Any], path: str) -> Any:
    current: Any = data
    for part in path.split("."):
        if isinstance(current, dict):
            current = current.get(part)
        else:
            return None
    return current


def replace_placeholders_in_text(text: str, mapping: Dict[str, str], data: Dict[str, Any], document_type: str) -> str:
    def replacer(match):
        raw_key = match.group(1)
        key = normalize_placeholder_name(raw_key, document_type)

        if key not in mapping:
            return ""

        value = get_value_from_path(data, mapping[key])
        return "" if value is None else str(value)

    return PLACEHOLDER_PATTERN.sub(replacer, text)


def _replace_text_in_paragraph(paragraph: Paragraph, mapping: Dict[str, str], data: Dict[str, Any], document_type: str) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    if not PLACEHOLDER_PATTERN.search(full_text):
        return

    new_text = replace_placeholders_in_text(full_text, mapping, data, document_type)
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _remove_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def cleanup_empty_paragraphs(doc: Document) -> None:
    for para in list(doc.paragraphs):
        if para.text.strip():
            continue
        parent_tag = para._element.getparent().tag.lower()
        if parent_tag.endswith("tc"):
            continue
        _remove_paragraph(para)


def render_docx_template(
    template_path: str | Path,
    output_path: str | Path,
    mapping: Dict[str, str],
    data: Dict[str, Any],
    document_type: str = "invoice",
    cleanup_empty_blocks: bool = True,
) -> Path:
    doc = Document(str(template_path))
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    for para in doc.paragraphs:
        _replace_text_in_paragraph(para, mapping, data, document_type)

    for table in doc.tables:
        rows_to_remove = []

        for row in list(table.rows):
            row_text = " ".join(cell.text for cell in row.cells)
            normalized_row_placeholders = {
                normalize_placeholder_name(p, document_type)
                for p in PLACEHOLDER_PATTERN.findall(row_text)
            }

            is_line_row = "LineDescription" in normalized_row_placeholders

            if is_line_row:
                for line in data.get("lines", []):
                    new_row = table.add_row()
                    for idx, cell in enumerate(row.cells):
                        raw_text = cell.text
                        raw_text = raw_text.replace("[[ItemDescription]]", "[[LineDescription]]")
                        raw_text = raw_text.replace("[[ItemQty]]", "[[LineQty]]")
                        raw_text = raw_text.replace("[[ItemUnitPrice]]", "[[LineUnitPrice]]")
                        raw_text = raw_text.replace("[[ItemTotal]]", "[[LineTotal]]")

                        raw_text = raw_text.replace("[[LineDescription]]", str(line.get("description", "")))
                        raw_text = raw_text.replace("[[LineQty]]", str(line.get("qty", "")))
                        raw_text = raw_text.replace("[[LineUnitPrice]]", str(line.get("unit_price", "")))
                        raw_text = raw_text.replace("[[LineTotal]]", str(line.get("total", "")))
                        raw_text = replace_placeholders_in_text(raw_text, mapping, data, document_type)
                        new_row.cells[idx].text = raw_text

                rows_to_remove.append(row)
            else:
                for cell in row.cells:
                    cell.text = replace_placeholders_in_text(cell.text, mapping, data, document_type)

        for row in rows_to_remove:
            table._tbl.remove(row._tr)

    if cleanup_empty_blocks:
        cleanup_empty_paragraphs(doc)

    doc.save(str(output_path))
    return output_path


def convert_docx_to_pdf(input_docx: str | Path, output_pdf: str | Path) -> Path:
    input_docx = Path(input_docx).resolve()
    output_pdf = Path(output_pdf).resolve()
    output_pdf.parent.mkdir(parents=True, exist_ok=True)

    system = platform.system().lower()

    if system == "windows":
        try:
            from docx2pdf import convert as docx2pdf_convert
        except ImportError as exc:
            raise RuntimeError("docx2pdf is not installed. Run: pip install docx2pdf") from exc

        docx2pdf_convert(str(input_docx), str(output_pdf))

        if not output_pdf.exists():
            raise RuntimeError("docx2pdf did not produce the PDF file.")

        return output_pdf

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "LibreOffice not installed.\n"
            "Linux: apt-get update -qq && apt-get install -y libreoffice\n"
            "Windows: pip install docx2pdf"
        )

    cmd = [
        soffice,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(output_pdf.parent),
        str(input_docx),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        raise RuntimeError(
            "PDF conversion failed.\n"
            f"STDOUT:\n{result.stdout}\n"
            f"STDERR:\n{result.stderr}"
        )

    generated_pdf = output_pdf.parent / f"{input_docx.stem}.pdf"
    if not generated_pdf.exists():
        raise RuntimeError("PDF was not generated.")

    if generated_pdf != output_pdf:
        generated_pdf.replace(output_pdf)

    return output_pdf


def _delete_with_retries(path: Path, retries: int = 10, delay: float = 0.4) -> None:
    for _ in range(retries):
        if not path.exists():
            return
        try:
            path.unlink()
            return
        except PermissionError:
            time.sleep(delay)


def render_mapped_template_to_pdf(
    template_path: str | Path,
    output_pdf_path: str | Path,
    standard_data: Dict[str, Any],
    mapping: Dict[str, str],
    document_type: str = "invoice",
    cleanup_empty_blocks: bool = True,
) -> Dict[str, Any]:
    placeholders = extract_placeholders_from_docx(template_path)

    normalized_mapping: Dict[str, str] = {}
    invalid_mapping_targets: Dict[str, str] = {}

    schema_paths = set(get_document_schema(document_type)["suggested_mapping"].values())

    for raw_key, target in mapping.items():
        normalized_key = normalize_placeholder_name(raw_key, document_type)
        normalized_mapping[normalized_key] = target

        if target not in schema_paths and not target.startswith("custom."):
            invalid_mapping_targets[raw_key] = target

    unmapped_fields = []
    for placeholder in placeholders:
        normalized = normalize_placeholder_name(placeholder, document_type)
        if normalized not in normalized_mapping:
            unmapped_fields.append(placeholder)

    output_pdf_path = Path(output_pdf_path)
    output_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="sales_doc_") as temp_dir:
        temp_docx_path = Path(temp_dir) / "rendered.docx"

        render_docx_template(
            template_path=template_path,
            output_path=temp_docx_path,
            mapping=normalized_mapping,
            data=standard_data,
            document_type=document_type,
            cleanup_empty_blocks=cleanup_empty_blocks,
        )

        convert_docx_to_pdf(temp_docx_path, output_pdf_path)
        _delete_with_retries(temp_docx_path)

    return {
        "ok": True,
        "document_type": document_type,
        "pdf_path": str(output_pdf_path),
        "unmapped_fields": unmapped_fields,
        "invalid_mapping_targets": invalid_mapping_targets,
        "detected_placeholders": placeholders,
    }
